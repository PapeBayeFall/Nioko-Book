import streamlit as st
import pandas as pd
import json
import os
import tempfile
import uuid
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# Configuration
st.set_page_config(layout="wide")
st.title("🛠️ Editeur de projet : Texte, Images, Tableaux et Export Word")

# Chemin de sauvegarde
SAVE_DIR = "temp_sauvegarde"
if not os.path.exists(SAVE_DIR):
    os.makedirs(SAVE_DIR)
SAVE_PATH = os.path.join(SAVE_DIR, "contenu.json")
IMAGES_DIR = os.path.join(SAVE_DIR, "images")
if not os.path.exists(IMAGES_DIR):
    os.makedirs(IMAGES_DIR)

# 🔵 Charger les données existantes
if "elements" not in st.session_state:
    if os.path.exists(SAVE_PATH):
        with open(SAVE_PATH, "r", encoding="utf-8") as f:
            st.session_state.elements = json.load(f)
    else:
        st.session_state.elements = []

# 🟢 Fonction pour sauver à chaque changement
def save_elements():
    with open(SAVE_PATH, "w", encoding="utf-8") as f:
        json.dump(st.session_state.elements, f, ensure_ascii=False, indent=4)

# 🔵 Fonctions d'ajout
def ajouter_paragraphe():
    titre = st.text_input("📝 Titre du paragraphe", key="titre_paragraphe")
    texte = st.text_area("✍️ Entrez votre texte", key="texte_paragraphe")
    if st.button("Ajouter ce paragraphe"):
        st.session_state.elements.append({
            "type": "paragraphe",
            "titre": titre,
            "contenu": texte,
            "media": [],
            "visible": True
        })
        save_elements()
        st.success("Paragraphe ajouté ✅")

def ajouter_image():
    if not st.session_state.elements:
        st.warning("Ajoutez d'abord un paragraphe pour y insérer une image.")
        return
    img = st.file_uploader("🖼️ Chargez une image", type=["png", "jpg", "jpeg"], key="img")
    titre_image = st.text_input("Titre pour l'image (optionnel, sera en italique)", key="titre_image")
    idx = st.selectbox("Sélectionnez un paragraphe pour l'insérer :", options=list(range(len(st.session_state.elements))), format_func=lambda i: st.session_state.elements[i]["titre"] or f"Paragraphe {i+1}")
    if img and st.button("Ajouter cette image au paragraphe sélectionné"):
        img_id = str(uuid.uuid4()) + ".png"
        img_path = os.path.join(IMAGES_DIR, img_id)
        with open(img_path, "wb") as f:
            f.write(img.read())
        st.session_state.elements[idx]["media"].append({
            "type": "image",
            "path": img_path,
            "titre": titre_image
        })
        save_elements()
        st.success("Image ajoutée ✅")



def ajouter_tableau():
    if not st.session_state.elements:
        st.warning("Ajoutez d'abord un paragraphe pour y insérer un tableau.")
        return

    fichier = st.file_uploader("📄 Chargez un fichier CSV ou Excel", type=["csv", "xlsx"], key="tableau")
    titre_tableau = st.text_input("Titre pour le tableau (optionnel, sera en italique)", key="titre_tableau")
    idx = st.selectbox(
        "Sélectionnez un paragraphe pour insérer le tableau :",
        options=list(range(len(st.session_state.elements))),
        format_func=lambda i: st.session_state.elements[i]["titre"] or f"Paragraphe {i+1}"
    )

    if fichier:
        try:
            if fichier.name.endswith('.csv'):
                df = pd.read_csv(fichier)
            else:
                df = pd.read_excel(fichier)

            # Supprimer l'index pour l'affichage
            df_reset = df.reset_index(drop=True)

            # Afficher le titre centré et en italique
            if titre_tableau:
                st.markdown(
                    f"<div style='text-align: center; font-style: italic; font-size: 20px;'>{titre_tableau}</div>",
                    unsafe_allow_html=True
                )
                st.write("")  # Petit espace

            # Afficher le tableau sans index
            st.dataframe(df_reset)

            if st.button("Ajouter ce tableau au paragraphe sélectionné"):
                df_json = df_reset.to_json(orient="split")
                st.session_state.elements[idx]["media"].append({
                    "type": "tableau",
                    "data": df_json,
                    "titre": titre_tableau
                })
                save_elements()
                st.success("Tableau ajouté ✅")
        except Exception as e:
            st.error(f"Erreur lors du chargement : {e}")


# 🔵 Sidebar
st.sidebar.header("➕ Ajouter un élément")
choix = st.sidebar.radio("Que voulez-vous ajouter ?", ("Paragraphe", "Image", "Tableau"))

if choix == "Paragraphe":
    ajouter_paragraphe()
elif choix == "Image":
    ajouter_image()
elif choix == "Tableau":
    ajouter_tableau()

# 🟢 Affichage des contenus
st.subheader("📜 Gestion du contenu actuel")

elements_to_delete = []

for idx, elem in enumerate(st.session_state.elements):
    with st.expander(f"{elem['titre'] or f'Paragraphe {idx+1}'}", expanded=True):
        if elem["visible"]:
            new_titre = st.text_input(
                "✏️ Modifier le titre",
                value=elem["titre"],
                key=f"title_edit_{idx}",
                on_change=lambda i=idx: update_titre(i)
            )

            texte = elem["contenu"]
            nb_lignes = texte.count('\n') + texte.count('.') // 3 + 5
            hauteur = max(100, nb_lignes * 20)

            st.text_area(
                "🖋️ Modifier le contenu", 
                value=texte, 
                key=f"edit_{idx}", 
                height=hauteur,
                on_change=lambda i=idx: update_contenu(i)
            )

            medias = elem.get("media", [])
            if medias:
                if len(medias) > 1:
                    cols = st.columns(2)
                else:
                    cols = [st.container()]

                for media_idx, media in enumerate(medias):
                    with cols[media_idx % len(cols)]:
                        if media["type"] == "image":
                            st.image(media["path"], use_container_width=True)
                            if media.get("titre"):
                                st.caption(f"_{media['titre']}_")
                        elif media["type"] == "tableau":
                            df = pd.read_json(media["data"], orient="split")
                            st.dataframe(df, use_container_width=False)
                            if media.get("titre"):
                                st.caption(f"_{media['titre']}_")

                        if st.button("❌ Supprimer ce média", key=f"delete_media_{idx}_{media_idx}", use_container_width=True):
                            del st.session_state.elements[idx]["media"][media_idx]
                            save_elements()
                            st.rerun()

        with st.container():
            col1, col2 = st.columns([5, 1])

            with col1:
                pass

            with col2:
                action = st.selectbox(
                    "⚙️ Option",
                    ["Aucune action", "Monter", "Descendre", "Masquer", "Afficher", "Supprimer"],
                    key=f"option_{idx}"
                )

                if action != "Aucune action":
                    if action == "Monter" and idx > 0:
                        st.session_state.elements[idx-1], st.session_state.elements[idx] = st.session_state.elements[idx], st.session_state.elements[idx-1]
                        save_elements()
                        st.rerun()
                    elif action == "Descendre" and idx < len(st.session_state.elements) - 1:
                        st.session_state.elements[idx], st.session_state.elements[idx+1] = st.session_state.elements[idx+1], st.session_state.elements[idx]
                        save_elements()
                        st.rerun()
                    elif action == "Masquer":
                        st.session_state.elements[idx]["visible"] = False
                        save_elements()
                        st.rerun()
                    elif action == "Afficher":
                        st.session_state.elements[idx]["visible"] = True
                        save_elements()
                        st.rerun()
                    elif action == "Supprimer":
                        elements_to_delete.append(idx)

# Suppression des éléments sélectionnés
if elements_to_delete:
    for idx in sorted(elements_to_delete, reverse=True):
        del st.session_state.elements[idx]
    save_elements()
    st.rerun()

def update_contenu(idx):
    st.session_state.elements[idx]["contenu"] = st.session_state[f"edit_{idx}"]
    save_elements()

def update_titre(idx):
    st.session_state.elements[idx]["titre"] = st.session_state[f"title_edit_{idx}"]
    save_elements()

# 📤 Génération du document Word
def creer_document(elements):
    doc = Document()
    doc.add_heading("Projet - Consulting Énergie", 0)

    for elem in elements:
        if elem.get("visible", True):
            if elem["titre"]:
                doc.add_heading(elem["titre"], level=1)
            p = doc.add_paragraph(elem["contenu"])
            p.style.font.name = 'Times New Roman'
            p.style.font.size = Pt(12)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            for media in elem.get("media", []):
                if media["type"] == "image":
                    doc.add_picture(media["path"], width=Inches(5))
                    if media.get("titre"):
                        p = doc.add_paragraph()
                        run = p.add_run(media["titre"])
                        run.italic = True
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                elif media["type"] == "tableau":
                    df = pd.read_json(media["data"], orient="split")
                    table = doc.add_table(rows=1, cols=len(df.columns))
                    hdr_cells = table.rows[0].cells
                    for i, col_name in enumerate(df.columns):
                        hdr_cells[i].text = str(col_name)
                    for index, row in df.iterrows():
                        row_cells = table.add_row().cells
                        for i, value in enumerate(row):
                            row_cells[i].text = str(value)
                    if media.get("titre"):
                        p = doc.add_paragraph()
                        run = p.add_run(media["titre"])
                        run.italic = True
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    return doc

# 📥 Bouton d'export
if st.button("📥 Générer le document Word"):
    doc = creer_document(st.session_state.elements)
    buffer = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(buffer.name)
    with open(buffer.name, "rb") as f:
        st.download_button("Télécharger le document Word", f, file_name="projet_consulting.docx")
    os.unlink(buffer.name)
